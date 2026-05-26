#!/usr/bin/env python3
"""Generate Aurex Hebrew RTL class/module documentation as DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(10)

# ── RTL helpers ────────────────────────────────────────────────────────────────

def _rtl(para, right_cm=0.0, hanging_cm=0.0):
    """Apply RTL direction + right-alignment + optional indentation."""
    pPr = para._element.get_or_add_pPr()
    # RTL paragraph direction
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    # Right alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)
    # Indentation
    if right_cm or hanging_cm:
        ind = OxmlElement('w:ind')
        if right_cm:
            ind.set(qn('w:right'), str(int(right_cm * 567)))
        if hanging_cm:
            ind.set(qn('w:hanging'), str(int(hanging_cm * 567)))
        pPr.append(ind)
    # Space after
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:after'), '40')
    pPr.append(spc)


def _heading(text, level=1):
    p = doc.add_heading(text, level)
    _rtl(p)
    return p


BULLETS = ['•', '○', '▪']
RIGHTS  = [0.0, 0.9, 1.8]  # cm from right margin for each level


def _b(level, text, bold=False):
    p = doc.add_paragraph()
    _rtl(p, right_cm=RIGHTS[level], hanging_cm=0.5)
    ch = BULLETS[min(level, 2)]
    run = p.add_run(ch + '  ')
    run.bold = False
    if bold:
        r2 = p.add_run(text)
        r2.bold = True
    else:
        p.add_run(text)
    return p


def _b_kv(level, key, val='', key_bold=True):
    """Bullet with bold key — plain value."""
    p = doc.add_paragraph()
    _rtl(p, right_cm=RIGHTS[level], hanging_cm=0.5)
    p.add_run(BULLETS[min(level, 2)] + '  ')
    r = p.add_run(key)
    r.bold = key_bold
    if val:
        p.add_run(' — ' + val)
    return p


def _method_bullet(name, params, ret):
    """Level-2 bullet: bold name(params) — ret."""
    p = doc.add_paragraph()
    _rtl(p, right_cm=RIGHTS[2], hanging_cm=0.5)
    p.add_run('▪  ')
    r = p.add_run(f'{name}({params})')
    r.bold = True
    if ret:
        p.add_run(f' — {ret}')
    return p


def spacer():
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:after'), '60')
    pPr.append(spc)


def add_class(name, role, attrs, methods):
    """Add one class documentation block."""
    _b_kv(0, name, role)
    if attrs:
        _b(1, 'תכונות:')
        for a in attrs:
            _b_kv(2, a[0], a[1] if len(a) > 1 else '')
    if methods:
        _b(1, 'פעולות:')
        for m in methods:
            _method_bullet(
                m[0],
                m[1] if len(m) > 1 else '',
                m[2] if len(m) > 2 else '',
            )


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
_heading('תיעוד מודולים ומחלקות — פרויקט Aurex', 0)

# ══════════════════════════════════════════════════════════════════════════════
# א. IMPORTED MODULES (TABLE)
# ══════════════════════════════════════════════════════════════════════════════
_heading('א. מודולים/מחלקות מיובאים', 1)

IMPORTS = [
    ('flet',               'ממשק משתמש גרפי מבוסס Flutter — חלונות, כפתורים, ניתוב דפים'),
    ('cryptography',       'הצפנה: RSA-OAEP, AES-CBC, ECDSA secp256k1, SHA-256, PKCS7'),
    ('threading',          'ניהול Thread-ים מקבילים, Lock, RLock, Event'),
    ('socket',             'תקשורת TCP ו-UDP ברמת מערכת ההפעלה'),
    ('json',               'סריאליזציה/דה-סריאליזציה של הודעות ונתוני DB ל-JSON'),
    ('pathlib',            'ניהול נתיבי קבצים ותיקיות חוצה-פלטפורמות'),
    ('hashlib',            'גיבוב SHA-256 לסיסמאות, תוכן קבצים ו-PoW'),
    ('base64',             'קידוד/פיענוח Base64 להעברת תמונות בינאריות ב-JSON'),
    ('queue',              'תורי הודעות thread-safe לתקשורת בין Thread-ים'),
    ('struct',             'אריזה/פירוק שדות אורך בינאריים בפרוטוקול הרשת'),
    ('dataclasses',        'הגדרת מחלקות נתונים עם __init__ ו-__repr__ אוטומטיים'),
    ('smtplib / ssl',      'שליחת אימיילי OTP דרך Gmail SMTP עם הצפנת TLS'),
    ('logging',            'מערכת לוגים מובנית של Python'),
    ('argparse',           'עיבוד ארגומנטים של שורת הפקודה'),
    ('os / sys',           'גישה לממשק מערכת ההפעלה, נתיבים ואינטרפטר Python'),
    ('random',             'יצירת מספרים פסאודו-אקראיים לקודי OTP ו-salt'),
    ('uuid',               'יצירת מזהי העלאה (upload_id) ייחודיים'),
    ('time / datetime',    'ניהול חותמות זמן, תוקף OTP ו-timestamp לבלוקים'),
    ('shutil',             'מחיקה רקורסיבית של תיקיות משתמש'),
    ('re',                 'ביטויים רגולריים לאימות שם משתמש ואימייל'),
    ('pickle',             'מיגרציה חד-פעמית של נתוני marketplace ישנים'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.columns[0].width = Cm(4)
tbl.columns[1].width = Cm(12)

hdr = tbl.rows[0].cells
hdr[0].text = 'שם המודול'
hdr[1].text = 'מטרה'
for cell in hdr:
    for para in cell.paragraphs:
        _rtl(para)
        for run in para.runs:
            run.bold = True

for mod, purpose in IMPORTS:
    row = tbl.add_row().cells
    row[0].text = mod
    row[1].text = purpose
    for cell in row:
        for para in cell.paragraphs:
            _rtl(para)

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# ב. DEVELOPED CLASSES
# ══════════════════════════════════════════════════════════════════════════════
_heading('ב. מחלקות שפותחו', 1)

# ─── SharedResources/classes.py ───────────────────────────────────────────────
_heading('SharedResources/classes.py', 2)

add_class(
    'Communication',
    'ניהול תקשורת מוצפנת AES-CBC בין כל רכיבי המערכת; תומך ב-async (תורים) וב-sync',
    [
        ('sock',               'סוקט TCP של החיבור הפעיל'),
        ('AES_key',            'מפתח AES-256 לסשן הנוכחי (מוחלף ב-RSA handshake)'),
        ('lock',               'Threading Lock להגנה על שליחה מקבילה'),
        ('msg_queue',          'תור הודעות נכנסות במצב async'),
        ('send_queue',         'תור הודעות יוצאות במצב async'),
        ('user',               'אובייקט המשתמש המחובר (בשימוש בצד שרת)'),
        ('peer_label',         'תווית הצד השני להצגה ב-logging'),
        ('_default_encryption','האם להצפין הודעות כברירת מחדל'),
        ('_close_marker',      'sentinel object לסיום תור ההודעות'),
    ],
    [
        ('send_one_message',        'data: dict, encryption=True',         'מצפין ושולח dict אחד (IV אקראי + AES-CBC + struct length prefix)'),
        ('recv_one_message',        'encryption=True',                      'מקבל, מפענח ומחזיר dict; מחזיר None אם החיבור נסגר'),
        ('AES_encrypt',             'plaintext, key, iv',                   'מצפין bytes ב-AES-CBC עם PKCS7 padding; מחזיר ciphertext bytes'),
        ('AES_decrypt',             'ciphertext, key, iv',                  'מפענח AES-CBC ומסיר padding; מחזיר plaintext bytes'),
        ('start_async',             'default_encryption=True',              'מפעיל Thread-י recv ו-send; הודעות מנותבות לתורים'),
        ('send_async',              'data: dict, encryption=None',          'מוסיף הודעה לתור השליחה; שולח ישירות אם async לא רץ'),
        ('recv_async',              'timeout=None',                         'שולף הודעה מתור הנכנסות; מחזיר None בסיום timeout'),
        ('is_close_marker',         'value',                                'מחזיר True אם value הוא ה-sentinel של סגירה'),
        ('stop_async',              '',                                      'מסיים Thread-י async ומוסיף None לתור השליחה'),
        ('set_user',                'user',                                 'מגדיר אובייקט משתמש על ה-comm (בשרת)'),
        ('connect',                 'ip, port',                             'מתחבר לכתובת TCP'),
        ('generate_iv',             '',                                     'מחזיר 16 בתים אקראיים ל-CBC IV (static)'),
        ('generate_AES_key',        '',                                     'מחזיר 16 בתים אקראיים כמפתח AES (static)'),
        ('close',                   '',                                     'עוצר async ומסגיר את הסוקט'),
        ('log',                     'dirct, data',                          'מדפיס שורת log בפורמט Recv From / Sent to'),
    ]
)

add_class(
    'RSA_Client',
    'לקוח TCP המבצע לחיצת יד RSA ומחזיק Communication מוצפן לתקשורת עם שרת',
    [
        ('ip / port',        'כתובת השרת'),
        ('sock',             'סוקט TCP'),
        ('communication',    'אובייקט Communication עם AES_key מוגדר'),
        ('RSA_public_key',   'מפתח RSA ציבורי של השרת (מתקבל ב-handshake)'),
        ('AES_key',          'מפתח AES שנוצר ונשלח מוצפן לשרת'),
    ],
    [
        ('start',                          '',       'מתחבר, מבצע RSA handshake, ומפעיל communicate_with_server'),
        ('contact_with_RSA',               '',       'לחיצת יד RSA: מקבל מפתח שרת, שולח AES מוצפן ב-RSA'),
        ('encrypt_AES_key_by_RSA_public_key', '',    'מצפין AES_key ב-RSA OAEP; מחזיר bytes מוצפנים'),
        ('communicate_with_server',        '',       'לולאת תקשורת ברירת מחדל (מיועד לדריסה)'),
        ('get_AES_key',                    '',       'מחזיר את מפתח ה-AES הנוכחי'),
        ('close',                          '',       'סוגר את הסוקט'),
    ]
)

add_class(
    'RSA_Server',
    'שרת TCP המאזין לחיבורים, מבצע לחיצת יד RSA לכל לקוח ומפעיל handle_client ב-Thread נפרד',
    [
        ('ip / port',      'כתובת האזנה'),
        ('name',           'שם הרכיב ל-logging'),
        ('peer_label',     'תווית הצד המחובר ל-logging'),
        ('dir_for_keys',   'תיקייה לשמירה/טעינה של מפתחות RSA PEM'),
        ('sock',           'סוקט TCP מאזין'),
        ('RSA_private_key','מפתח RSA פרטי 2048-bit'),
        ('RSA_public_key', 'מפתח RSA ציבורי'),
    ],
    [
        ('start',                  '',                      'לולאת accept; מפעיל Thread ל-communicate_with_client לכל חיבור'),
        ('communicate_with_client','client_socket',         'מבצע RSA handshake ומפעיל handle_client עם Communication'),
        ('handle_client',          'communication',         'מיועד לדריסה; עיבוד הודעות לאחר handshake'),
        ('contact_with_RSA',       'communication',         'מבצע לחיצת יד RSA: שולח מפתח ציבורי, מפענח AES'),
        ('create_RSA_keys',        'dir_for_keys',          'טוען או מייצר מפתחות RSA ושומר ל-PEM'),
        ('get_public_key_RSA',     '',                      'מחזיר מפתח RSA ציבורי כ-bytes PEM'),
        ('get_encrypted_AES_key',  'data: bytes, communication', 'מפענח AES_key עם RSA פרטי ושומר ב-communication.AES_key'),
    ]
)

add_class(
    'UDPServer',
    'שרת UDP המגיב לשאילתות גילוי WHRSV מ-Bnode-ים עם כתובת ה-Gateway',
    [
        ('sock',            'סוקט UDP'),
        ('ip / port',       'כתובת האזנה UDP'),
        ('srv_ip / srv_port','כתובת שרת ה-Gateway להחזיר בתגובה'),
        ('message_to_send', 'מחרוזת תגובה: SRVAT|ip|port (מקודדת ל-bytes)'),
    ],
    [
        ('run', '', 'לולאה: מקבל WHRSV מ-Bnode, שולח בחזרה SRVAT|ip|port'),
    ]
)

add_class(
    'UDPClient',
    'לקוח UDP לשידור broadcast לגילוי כתובת TCP של ה-Gateway',
    [
        ('sock',          'סוקט UDP עם SO_BROADCAST'),
        ('udp_srv_port',  'פורט UDP של ה-Gateway'),
        ('tcp_ip / tcp_port', 'כתובת TCP שהתקבלה לאחר גילוי'),
    ],
    [
        ('run', '', 'שולח WHRSV ב-broadcast; ממתין ל-SRVAT|ip|port; מחזיר (ip, port)'),
    ]
)

add_class(
    'User (classes.py)',
    'מודל משתמש עם גיבוב סיסמה PEPPER+salt+SHA-256, קוד OTP ויתרת ארנק (גרסת classes.py)',
    [
        ('username',           'שם המשתמש'),
        ('email',              'כתובת אימייל'),
        ('salt',               'ערך salt אקראי לגיבוב'),
        ('password_hash',      'SHA-256(PEPPER + password + salt)'),
        ('pubkey',             'מפתח ECDSA ציבורי'),
        ('verification_code',  'קוד OTP לאיפוס סיסמה'),
        ('reset_time',         'ISO timestamp לפקיעת OTP'),
        ('wallet_balance',     'יתרת ארנק (float)'),
    ],
    [
        ('verify_password',            'password',              'מחזיר True אם hash(password) תואם password_hash'),
        ('set_verification_code',      'code',                  'מגדיר verification_code'),
        ('set_reset_time',             'time',                  'מגדיר reset_time'),
        ('is_code_match_and_available','current_time, code',    'מחזיר True אם הקוד תואם ולא פג תוקפו'),
        ('set_password',               'new_password',          'מעדכן password_hash לפי הסיסמה החדשה'),
    ]
)

add_class(
    'Transaction (dataclass)',
    'ייצוג עסקה בלוקצ\'יין: שולח, מקבל, סכום וחתימה דיגיטלית',
    [
        ('sender',    'מפתח ציבורי של השולח'),
        ('receiver',  'מפתח ציבורי של המקבל'),
        ('amount',    'סכום AUR להעברה'),
        ('signature', 'חתימה דיגיטלית ECDSA'),
        ('timestamp', 'חותמת זמן UNIX'),
    ],
    []
)

add_class(
    'Block (dataclass)',
    'בלוק בלוקצ\'יין: מכיל עסקה, prev_hash, nonce ותאריך',
    [
        ('index',       'אינדקס הבלוק בשרשרת'),
        ('prev_hash',   'גיבוב הבלוק הקודם'),
        ('transaction', 'אובייקט Transaction'),
        ('nonce',       'מונה PoW'),
        ('timestamp',   'חותמת זמן יצירה'),
    ],
    [
        ('compute_hash', '', 'מחשב SHA-256 של ה-dict של הבלוק; מחזיר hex string'),
    ]
)

spacer()

# ─── SharedResources/logging.py ───────────────────────────────────────────────
_heading('SharedResources/logging.py', 2)

add_class(
    'Logger',
    'מנהל לוגים צבעוני: כתום ל-WARNING, אדום ל-ERROR, צהוב לכתובות; מסנן רעשי Flet',
    [
        ('file_name',    'שם הקובץ שיצר את ה-Logger'),
        ('_logger',      'אובייקט logging.Logger הפנימי'),
        ('_level',       'רמת לוג גלובלית (class variable)'),
    ],
    [
        ('debug',     'message',     'כותב הודעת DEBUG'),
        ('info',      'message',     'כותב הודעת INFO'),
        ('warning',   'message',     'כותב הודעת WARNING בצבע כתום'),
        ('error',     'message',     'כותב הודעת ERROR בצבע אדום'),
        ('set_level', 'level_str',   'מגדיר רמת לוג גלובלית ("DEBUG"/"INFO"/"WARNING"/"ERROR") — class method'),
    ]
)

spacer()

# ─── Server/DB_ORM.py ─────────────────────────────────────────────────────────
_heading('Server/DB_ORM.py', 2)

add_class(
    'User (DB_ORM)',
    'מודל משתמש לאחסון ב-DB: גיבוב סיסמה PEPPER+salt, מפתח ECDSA, קוד OTP',
    [
        ('username',          'שם משתמש ייחודי (stripped)'),
        ('email',             'אימייל (lowercase)'),
        ('salt',              'ערך salt אקראי'),
        ('password_hash',     'SHA-256(PEPPER + password + salt)'),
        ('public_key',        'מפתח ECDSA secp256k1 ציבורי (hex uncompressed)'),
        ('verification_code', 'קוד OTP לאיפוס סיסמה'),
        ('reset_time',        'ISO timestamp פקיעת OTP'),
    ],
    [
        ('verify_password',             'password',           'מחזיר True אם הסיסמה תואמת ה-hash'),
        ('set_password',                'new_password',       'מחשב ומעדכן password_hash'),
        ('set_public_key',              'public_key',         'מעדכן public_key'),
        ('set_verification_code',       'code',               'מגדיר verification_code'),
        ('set_reset_time',              'value',              'מגדיר reset_time (ISO string)'),
        ('is_code_match_and_available', 'current_time, code','מחזיר True אם הקוד תואם ולא פג תוקפו'),
        ('to_dict',                     '',                   'מחזיר dict לשמירה ב-JSON'),
        ('from_dict',                   'raw: dict',          'מייצר User מ-dict; class method'),
    ]
)

add_class(
    'MarketplaceItem (DB_ORM)',
    'נכס דיגיטלי: מטא-דטה, מצב (PENDING/FOR_SALE/UNLISTED/SOLD), גרסה ומפתח ECDSA',
    [
        ('asset_id',     'מזהה נכס ייחודי (16 תווים hex)'),
        ('owner',        'שם משתמש הבעלים הנוכחי'),
        ('asset_name',   'שם הנכס'),
        ('description',  'תיאור חופשי'),
        ('file_type',    'jpg / png'),
        ('cost',         'מחיר ב-AUR'),
        ('storage_path', 'נתיב מלא לקובץ התמונה ב-DB/uploads/'),
        ('created_at',   'ISO timestamp יצירה'),
        ('version',      'מספר גרסה (עולה בכל שינוי בעלות/סטטוס)'),
        ('asset_status', 'PENDING / FOR_SALE / UNLISTED / SOLD'),
        ('public_key',   'מפתח ECDSA של הבעלים בעת העלאה'),
    ],
    [
        ('to_dict',   '',          'מחזיר dict לשמירה ב-JSON'),
        ('from_dict', 'raw: dict', 'מייצר MarketplaceItem מ-dict; מבצע מיגרציה של שדות ישנים; class method'),
    ]
)

add_class(
    'ORM',
    'ממשק גישה מאוחד ל-users.json, marketplace_items.json ו-notifications.json; thread-safe עם RLock',
    [
        ('users_json_path',         'נתיב ל-users.json'),
        ('marketplace_json_path',   'נתיב ל-marketplace_items.json'),
        ('notifications_json_path', 'נתיב ל-notifications.json'),
        ('_lock',                   'RLock לכתיבה בטוחה מ-Thread-ים מרובים'),
    ],
    [
        ('add_user',             'username, password, email',     'יוצר User חדש אם השם והאימייל פנויים; מחזיר (bool, message)'),
        ('get_user',             'username',                       'מחזיר User לפי שם, או None'),
        ('save_user',            'user: User',                    'שומר/מעדכן User ב-JSON'),
        ('get_user_by_email',    'email',                         'מחזיר User לפי אימייל, או None'),
        ('get_user_by_public_key','public_key',                   'מחזיר User לפי מפתח ECDSA ציבורי, או None'),
        ('set_user_public_key',  'username, public_key',          'מעדכן public_key של משתמש; מחזיר bool'),
        ('issue_reset_code',     'email, minutes_valid=5',        'מייצר OTP 6 ספרות עם תוקף; שומר; מחזיר (bool, msg, code)'),
        ('verify_reset_code',    'email, code',                   'מאמת OTP ותוקפו; מחזיר (bool, message, User|None)'),
        ('update_password_by_email','email, new_password',        'מחשב hash חדש ושומר; מחזיר (bool, message)'),
        ('delete_user',          'username',                      'מוחק משתמש; מחזיר bool'),
        ('delete_user_assets',   'username',                      'מוחק כל נכסי המשתמש מה-marketplace'),
        ('add_asset',            'username, asset: MarketplaceItem','מוסיף נכס לרשימת הבעלים; מחזיר True'),
        ('get_all_assets',       '',                              'מחזיר כל הנכסים ממוינים לפי created_at (חדש ראשון)'),
        ('get_all_for_sale_assets','',                            'מחזיר נכסים בסטטוס FOR_SALE בלבד, ממוינים'),
        ('update_asset_status',  'asset_id, status, increment_version=False','מעדכן asset_status (ו-version אם נדרש); מחזיר bool'),
        ('transfer_asset',       'asset_id, from_owner, to_owner','מעביר נכס בין בעלים: סטטוס→UNLISTED, גרסה+1; מחזיר bool'),
        ('delete_asset',         'asset_id, owner',              'מוחק נכס מה-DB; מחזיר bool'),
        ('get_assets_for_user',  'username',                      'מחזיר נכסי משתמש שאינם FOR_SALE (My Assets)'),
        ('find_asset_by_id',     'asset_id',                      'מחזיר MarketplaceItem לפי asset_id, או None'),
        ('queue_notification',   'username, msg',                 'מוסיף התראה לתור ה-JSON של המשתמש'),
        ('flush_notifications',  'username',                      'מחזיר ומוחק את כל ההתראות הממתינות של המשתמש'),
    ]
)

spacer()

# ─── Server/server_module.py ──────────────────────────────────────────────────
_heading('Server/server_module.py', 2)

add_class(
    'UploadSession (dataclass)',
    'מצב העלאת קובץ בין UPLOAD_INIT ל-UPLOAD_FINISH; מאוחסן בזיכרון בלבד',
    [
        ('upload_id',   'מזהה UUID ייחודי'),
        ('username',    'שם המשתמש המעלה'),
        ('asset_name',  'שם הנכס'),
        ('description', 'תיאור'),
        ('file_type',   'jpg / png'),
        ('cost',        'מחיר ב-AUR'),
        ('chunks_b64',  'רשימת חתיכות Base64 שהתקבלו מהלקוח'),
        ('created_at',  'ISO timestamp פתיחת הסשן'),
    ],
    []
)

add_class(
    'ORMExtended',
    'הרחבת ORM לצד שרת: מוסיף RLock לשמירה, lookup לפי אימייל, ניהול OTP ועדכון סיסמה עם לוג',
    [],
    [
        ('save_user',              'user',               'שומר User ב-JSON תחת RLock'),
        ('get_user_by_email',      'email',              'מחזיר User לפי אימייל'),
        ('set_user_public_key',    'username, public_key','מעדכן public_key ושומר'),
        ('issue_reset_code',       'email, minutes_valid=5','מייצר OTP עם תוקף ושומר; מחזיר (bool, msg, code)'),
        ('verify_reset_code',      'email, code',        'מאמת OTP ותוקפו; מחזיר (bool, message, User|None)'),
        ('update_password_by_email','email, new_password','מעדכן hash ומדפיס לוג השוואת hash; מחזיר (bool, msg)'),
    ]
)

add_class(
    'ServerUpdated',
    'שרת מרכזי: מנהל חיבורי לקוחות וGateway-ים, DB, push events ושיגור ל-blockchain',
    [
        ('host / port',        'כתובת השרת'),
        ('db',                 'אובייקט ORMExtended לגישה ל-DB'),
        ('upload_sessions',    'dict: upload_id → UploadSession פעיל'),
        ('gateway_clients',    'set של Communication-ים של Gateway-ים מחוברים'),
        ('online_users',       'dict: username → Communication של לקוח מחובר'),
        ('handlers',           'dict: type → פונקציה handler לניתוב הודעות'),
        ('client_listener',    'RSA_Server להאזנה ללקוחות'),
        ('notifications_path', 'נתיב notifications.json'),
    ],
    [
        ('start',                   '',                             'מפעיל client_listener'),
        ('handle_client',           'comm',                         'לולאת קריאה: מנתב ל-dispatch, שולח תגובה'),
        ('dispatch',                'comm, msg',                    'מחפש handler לפי msg["type"]; מחזיר dict תגובה'),
        ('handle_login',            'comm, msg',                    'מאמת username/password; מחזיר LOGIN_SUCCESS + username או ERROR'),
        ('handle_signup',           'comm, msg',                    'רושם חשבון חדש; מחזיר SIGNUP_SUCCESS או ERROR'),
        ('handle_send_code',        'comm, msg',                    'מייצר OTP ושולח לאימייל; מחזיר CODE_SENT'),
        ('handle_verify_code',      'comm, msg',                    'מאמת OTP; מחזיר CODE_VERIFIED + username'),
        ('handle_update_password',  'comm, msg',                    'מאמת OTP ומעדכן סיסמה; מחזיר PASSWORD_UPDATED'),
        ('handle_logout',           'comm, msg',                    'מסיר ממפת online_users; מחזיר LOGOUT_SUCCESS'),
        ('handle_upload_init',      'comm, msg',                    'יוצר UploadSession; מחזיר UPLOAD_READY'),
        ('handle_upload',           'comm, msg',                    'מוסיף chunk לסשן; מחזיר CHUNK_RECEIVED'),
        ('handle_upload_finish',    'comm, msg',                    'מחבר chunks, שומר קובץ, יוצר MarketplaceItem; מחזיר UPLOAD_SUCCESS + asset_id'),
        ('handle_move_to_marketplace','comm, msg',                  'שולח UPLOAD_ASSET ל-Gateway לכרייה; מחזיר MOVE_PENDING'),
        ('handle_update_public_key','comm, msg',                    'שומר מפתח ECDSA ושולח CREATE_BALANCE ל-Gateway; מחזיר KEY_UPDATED'),
        ('handle_register_gateway', 'comm, msg',                    'מוסיף comm ל-gateway_clients; מחזיר GATEWAY_REGISTERED'),
        ('handle_buy_asset',        'comm, msg',                    'מעביר TX_REQUEST_BUY ל-Gateway; מחזיר BUY_SUBMITTED'),
        ('handle_buy_success',      'comm, msg',                    'מעביר בעלות נכס, שולח push BUY_SUCCESS לקונה ולמוכר; מחזיר BUY_ACKNOWLEDGED'),
        ('handle_sell_success',     'comm, msg',                    'שולח push BLOCK_ACCEPTED למוכר; מחזיר SELL_ACKNOWLEDGED'),
        ('handle_block_rejected',   'comm, msg',                    'שולח push BLOCK_REJECTED למשתמש; מחזיר BLOCK_REJECTED_ACKNOWLEDGED'),
        ('handle_send_balance',     'comm, msg',                    'מקבל יתרה מ-Gateway, מאתר משתמש לפי public key, שולח push BALANCE_IS'),
        ('handle_get_balance',      'comm, msg',                    'שולח GET_BALANCE ל-Gateway; מחזיר BALANCE_REQUESTED'),
        ('handle_fully_upload',     'comm, msg',                    'מסמן FOR_SALE, שולח FULLY_UPLOADED ל-owner, מרענן יתרה; מחזיר FULLY_UPLOAD_ACKNOWLEDGED'),
        ('handle_asset_unlisted',   'comm, msg',                    'מסמן UNLISTED, שולח ASSET_UNLISTED ל-owner; מחזיר UNLIST_ACKNOWLEDGED'),
        ('handle_get_assets_ids',   'comm, msg',                    'מחזיר רשימת IDs: FOR_SALE (שוק) או נכסי משתמש (My Assets)'),
        ('handle_get_asset_by_id',  'comm, msg',                    'שולח streaming: ASSET_INIT + ASSET_CHUNK × n + ASSET_END'),
        ('handle_delete_account',   'comm, msg',                    'מוחק משתמש, נכסים, התראות ותיקיית uploads; מחזיר ACCOUNT_IS_DELETED'),
        ('_notify_gateways',        'payload',                      'שולח הודעה לכל gateway_clients המחוברים'),
        ('_push_event',             'username, event',              'שולח event ל-online user; אם offline — מכניס לתור ההתראות'),
        ('_queue_notification',     'username, msg',                'שומר התראת טקסט ב-notifications.json'),
        ('_flush_notifications_for_user','username, comm',          'שולח כל ההתראות הממתינות ללקוח עם התחברות'),
    ]
)

spacer()

# ─── Client/wallet_manager.py ─────────────────────────────────────────────────
_heading('Client/wallet_manager.py', 2)

add_class(
    'WalletData',
    'נתוני ארנק ECDSA secp256k1: מפתחות ציבורי/פרטי, חתימה ואימות חתימות',
    [
        ('username',    'שם המשתמש הבעלים'),
        ('public_key',  'מפתח ECDSA ציבורי hex (04…, 65 בתים uncompressed)'),
        ('private_key', 'מפתח ECDSA פרטי hex (32 בתים)'),
    ],
    [
        ('validate',          '',                              'מוודא שהמפתח הפרטי מייצר את הציבורי; מחזיר (bool, reason)'),
        ('sign_payload',      'payload: dict',                 'חותם canonical_json(payload) ב-ECDSA SHA-256; מחזיר signature hex'),
        ('verify_signature',  'payload: dict, signature_hex',  'מאמת חתימה ECDSA; מחזיר bool'),
        ('to_dict',           '',                              'מחזיר dict לשמירה ב-JSON'),
        ('from_dict',         'raw: dict',                     'מייצר WalletData מ-dict; class method'),
    ]
)

add_class(
    'WalletManager',
    'ניהול שמירה וטעינת ארנקים מקומיים תחת Client/<username>/wallet.json',
    [
        ('base_dir', 'תיקיית בסיס לארנקים (ברירת מחדל: Client/)'),
    ],
    [
        ('wallet_path_for_user', 'username',        'מחזיר Path: Client/<username>/wallet.json'),
        ('generate_wallet',      'username',        'מייצר זוג מפתחות ECDSA, שומר ומחזיר WalletData'),
        ('save_wallet',          'wallet, path',    'שומר WalletData כ-JSON לנתיב הנתון'),
        ('load_wallet_from_path','path',            'טוען, מאמת ומחזיר WalletData; זורק ValueError אם פגום'),
        ('load_wallet_for_user', 'username',        'טוען ארנק; בודק גם נתיב legacy; מחזיר WalletData או None'),
    ]
)

spacer()

# ─── Client/client.py ─────────────────────────────────────────────────────────
_heading('Client/client.py', 2)

add_class(
    'MarketItem (dataclass)',
    'פריט שוק עבור ממשק המשתמש — נגזר מהודעת ASSET_INIT שמגיעה מהשרת',
    [
        ('asset_id',       'מזהה נכס'),
        ('owner',          'שם בעלים'),
        ('title',          'שם הנכס'),
        ('description',    'תיאור'),
        ('file_type',      'jpg / png'),
        ('price',          'מחיר ב-AUR'),
        ('created_at',     'תאריך יצירה'),
        ('public_key_hex', 'מפתח ECDSA ציבורי של הבעלים'),
        ('asset_status',   'PENDING / FOR_SALE / UNLISTED'),
        ('version',        'גרסת הנכס'),
    ],
    []
)

add_class(
    'ImageCache',
    'מטמון תמונות לכל משתמש: RAM + דיסק ב-Client/<user>/cache/; מנהל מטא-דטה ויתרה ב-metadata.json',
    [
        ('username',        'שם המשתמש'),
        ('_ram',            'dict: asset_id → bytes תמונה בזיכרון'),
        ('_lock',           'Threading Lock לגישה בטוחה'),
        ('_cache_dir',      'נתיב תיקיית מטמון'),
        ('_assets_dir',     'נתיב תת-תיקייה לקבצי תמונה'),
        ('_metadata_path',  'נתיב metadata.json'),
        ('_metadata',       'dict: "balance"→float + asset_id→entry_dict'),
    ],
    [
        ('get_balance',     '',                                      'מחזיר float יתרה מתוך _metadata'),
        ('set_balance',     'amount: float',                         'מעדכן יתרה ב-_metadata ושומר ל-metadata.json'),
        ('get_raw',         'asset_id',                              'מחזיר bytes תמונה מ-RAM או דיסק, או None'),
        ('get_path',        'asset_id',                              'מחזיר Path לקובץ תמונה אם קיים, או None'),
        ('get_if_current',  'asset_id, server_version',              'מחזיר (entry_dict, bytes) אם גרסה מספיקה; אחרת None'),
        ('store',           'asset_id, file_type, version, meta, raw','שומר bytes + metadata לדיסק ול-RAM'),
        ('invalidate',      'asset_id',                              'מוחק נכס מ-RAM, metadata ודיסק'),
    ]
)

add_class(
    'ClientState (dataclass)',
    'מצב UI ופעלון ממשק המשתמש הנוכחי (in-memory בלבד)',
    [
        ('username',             'שם משתמש מחובר'),
        ('email',                'אימייל המשתמש'),
        ('is_authenticated',     'האם מחובר לשרת'),
        ('market_items',         'רשימת MarketItem-ים שנטענו'),
        ('notifications',        'רשימת הודעות'),
        ('unseen_notifications', 'מונה התראות שלא נקראו'),
        ('wallet_loaded',        'האם ארנק נטען'),
        ('wallet_public_key',    'מפתח ECDSA ציבורי הנוכחי'),
        ('balance',              'יתרת AUR הנוכחית (float)'),
    ],
    []
)

add_class(
    'ServerClient',
    'עטיפת תקשורת sync לפרוטוקול dict עם RSA_Client; מנתב push events לתורים ייעודיים',
    [
        ('_comm',               'אובייקט Communication הפעיל'),
        ('_response_queue',     'תור תגובות request-response'),
        ('notification_queue',  'תור הודעות NOTIFICATION / BUY / BLOCK'),
        ('balance_queue',       'תור עדכוני יתרה BALANCE_IS'),
        ('asset_sold_queue',    'תור asset_id-ים שנמכרו'),
        ('asset_unlisted_queue','תור asset_id-ים שהורדו מהשוק'),
    ],
    [
        ('connect',          '',                                    'מתחבר לשרת ומפעיל Thread קבלה'),
        ('close',            '',                                    'עוצר Thread ומנתק מהשרת'),
        ('login',            'username, password',                  'שולח LOGIN; מחזיר dict תגובה'),
        ('signup',           'username, password, email',           'שולח SIGNUP; מחזיר dict תגובה'),
        ('send_code',        'email',                              'שולח SEND_CODE; מחזיר CODE_SENT'),
        ('verify_code',      'email, code',                        'שולח VERIFY_CODE; מחזיר CODE_VERIFIED'),
        ('update_password',  'email, new_password, code',          'שולח UPDATE_PASSWORD; מחזיר PASSWORD_UPDATED'),
        ('upload_file',      'username, file_path, name, desc, type, cost, ...','שולח UPLOAD_INIT → UPLOAD-chunks → UPLOAD_FINISH; מחזיר UPLOAD_SUCCESS + asset_id'),
        ('download_asset',   'asset_id, timeout=30',               'מקבל ASSET_INIT + ASSET_CHUNK × n + ASSET_END; מחזיר (meta_dict, raw_bytes)'),
        ('move_to_marketplace','username, asset_id',               'שולח MOVE_TO_MARKETPLACE; מחזיר MOVE_PENDING'),
        ('buy_asset',        'payload',                            'שולח BUY_ASSET; מחזיר BUY_SUBMITTED'),
        ('unlist_asset',     'username, asset_id, public_key, sig','שולח UNLIST_ASSET; מחזיר תגובה משרת'),
        ('request_balance',  'public_key',                         'שולח GET_BALANCE; יתרה תגיע כ-BALANCE_IS push event'),
        ('get_assets_ids',   'username=""',                        'שולח GET_ASSETS_IDS; מחזיר ASSETS_IDS_LIST'),
        ('delete_account',   'username',                           'שולח DELETE_ACCOUNT; מחזיר ACCOUNT_IS_DELETED'),
        ('update_public_key','username, public_key',               'שולח UPDATE_PUBLIC_KEY; מחזיר KEY_UPDATED'),
    ]
)

add_class(
    'ClientApp',
    'בקר Flet הראשי: ניהול מצב, ניתוב דפים, wallet, cache ופעולות עסקיות מול השרת',
    [
        ('page',                'דף Flet'),
        ('client',              'אובייקט ServerClient'),
        ('wallet_manager',      'אובייקט WalletManager'),
        ('state',               'אובייקט ClientState'),
        ('wallet_session',      'WalletData הנוכחי (None אם לא טעון)'),
        ('_image_cache',        'ImageCache למשתמש המחובר הנוכחי'),
        ('_balance_text',       'ref ל-ft.Text להצגת יתרה ב-header'),
        ('_notification_badge', 'ref ל-ft.Container לאייקון התראות'),
    ],
    [
        ('start',                     '',                                      'מנווט ל-/login'),
        ('login',                     'username, password',                    'מאמת, טוען ארנק אוטומטי, מרענן התראות; מחזיר dict'),
        ('signup',                    'username, password, email',             'רושם חשבון חדש; מחזיר dict'),
        ('send_code',                 'email',                                 'מבקש OTP לאימייל; מחזיר dict'),
        ('verify_code',               'email, code',                           'מאמת OTP, שומר verified_reset_user; מחזיר dict'),
        ('update_password',           'email, new_password, code',             'מעדכן סיסמה לאחר OTP; מחזיר dict'),
        ('logout',                    '',                                      'מנתק, מאפס state, cache, wallet ו-refs לUI'),
        ('delete_account',            '',                                      'מוחק חשבון מהשרת ומאפס את כל המצב'),
        ('upload_asset',              'file_path, name, desc, type, cost, for_sale=True','מעלה נכס; אם for_sale=True קורא MOVE_TO_MARKETPLACE אוטומטי'),
        ('move_to_marketplace',       'asset_id',                              'שולח MOVE_TO_MARKETPLACE; מחזיר dict'),
        ('buy_asset',                 'item: MarketItem',                      'חותם ושולח BUY_ASSET עם חתימת ECDSA; מחזיר dict'),
        ('unlist_asset',              'asset_id',                              'חותם ושולח UNLIST_ASSET; מחזיר dict'),
        ('load_asset_by_id',          'asset_id, version=1',                   'טוען נכס מ-cache או מהשרת; מחזיר MarketItem או None'),
        ('get_market_asset_ids',      '',                                      'שולח GET_ASSETS_IDS ללא username; מחזיר [{id, version}]'),
        ('get_my_asset_ids',          '',                                      'שולח GET_ASSETS_IDS עם username; מחזיר [{id, version}]'),
        ('request_balance',           '',                                      'שולח GET_BALANCE עם public_key; יתרה תגיע כ-push event'),
        ('update_public_key',         'public_key',                            'מעדכן בשרת ומבקש רענון יתרה מ-blockchain'),
        ('sign_payload',              'payload: dict',                         'חותם dict עם WalletData הנוכחי; מחזיר signature hex'),
        ('generate_new_wallet',       '',                                      'מייצר ארנק ECDSA חדש, שומר ורושם בשרת'),
        ('load_wallet_from_file',     'file_path',                             'טוען ארנק מקובץ חיצוני, שומר כברירת מחדל'),
        ('load_default_wallet',       '',                                      'טוען ארנק מ-Client/<user>/wallet.json'),
        ('export_wallet',             'output_path',                           'מייצא ארנק נוכחי לנתיב חיצוני'),
        ('_drain_balance_events',     '',                                      'שולף עדכוני יתרה מ-balance_queue; מעדכן state, cache ו-UI'),
        ('_start_balance_monitor',    '',                                      'מפעיל Thread שרץ כל 0.5 שניות ל-drain_balance_events'),
        ('_start_notification_monitor','',                                     'מפעיל Thread שרץ כל 0.4 שניות לעדכון badge התראות'),
    ]
)

spacer()

# ─── Gateway/gateway.py ───────────────────────────────────────────────────────
_heading('Gateway/gateway.py', 2)

add_class(
    'GatewayServer',
    'שרת שער מרכזי בטופולוגיית כוכב: מנתב הודעות בין שרת ראשי לצמתי blockchain ומאמת בלוקים',
    [
        ('node_listener',         'RSA_Server להאזנה לצמתי blockchain'),
        ('server_client',         'RSA_Client לחיבור לשרת הראשי'),
        ('udp_service',           'UDPServer לתגובת גילוי WHRSV'),
        ('nodes',                 'dict: (ip, port) → {comm, chain_length, registered}'),
        ('gateway_operations',    'dict: type → handler להודעות מהשרת הראשי'),
        ('blockchain_operations', 'dict: type → handler להודעות מצמתי blockchain'),
    ],
    [
        ('start',                           '',             'מפעיל UDP, server_client ו-node_listener ב-Thread-ים'),
        ('stop',                            '',             'מסמן stop_event לעצירה'),
        ('handle_node_connection',          'comm',         'לולאת קריאה מ-Bnode; מנתב ל-blockchain_operations'),
        ('communicate_with_main_server',    '',             'שולח REGISTER_GATEWAY; לולאת קריאה מהשרת; מנתב ל-gateway_operations'),
        ('register_blockchain_node',        'request, comm','רושם Bnode בטבלת nodes עם chain_length'),
        ('tx_request_buy',                  'request',      'מעביר TX_REQUEST_BUY לכל הצמתים'),
        ('tx_request_sell',                 'request',      'מעביר TX_REQUEST_SELL לכל הצמתים'),
        ('broadcast_tx_to_verify',          'request, comm','מאמת בלוק (חתימה + PoW), מסנכרן צמתים מפגרים, מפיץ לשאר'),
        ('handle_get_balance',              'request',      'מפיץ GET_BALANCE לכל הצמתים'),
        ('notify_buy_success',              'request',      'מעביר BUY_SUCCESS לשרת הראשי'),
        ('notify_sell_success',             'request',      'מעביר SELL_SUCCESS לשרת הראשי'),
        ('notify_send_balance',             'request',      'מעביר SEND_BALANCE (עם יתרה) לשרת הראשי'),
        ('handle_asset_signed_in_blockchain','request',     'מאמת בלוק ASSET_MINT; שולח FULLY_UPLOAD לשרת'),
        ('handle_asset_unlist_signed_in_blockchain','request','מאמת בלוק UNLIST; שולח ASSET_UNLISTED לשרת'),
        ('create_balance',                  'request',      'מפיץ CREATE_BALANCE לכל הצמתים ליצירת יתרה ראשונית'),
        ('handle_unlist_to_nodes',          'request',      'מפיץ UNLIST_ASSET לכל הצמתים לכרייה'),
        ('upload_asset_to_nodes',           'request',      'מפיץ UPLOAD_ASSET לכל הצמתים לכרייה'),
        ('_validate_block',                 'block',        'מאמת חתימת ECDSA וגיבוב PoW; מחזיר (bool, reason)'),
        ('_broadcast_to_nodes',             'msg, skip_addr=None','שולח הודעה לכל הצמתים הפעילים (מדלג על skip_addr)'),
        ('_route_to_server',                'msg',          'שולח הודעה לשרת הראשי'),
    ]
)

spacer()

# ─── blockchain/Bnode.py ──────────────────────────────────────────────────────
_heading('blockchain/Bnode.py', 2)

add_class(
    'BlockchainNode',
    'צומת blockchain עצמאי: כרייה PoW, ניהול ledger ויתרות, גילוי Gateway ב-UDP, סנכרון עמיתים',
    [
        ('ip / port',        'כתובת האזנה לעמיתים (OS-assigned אם 0)'),
        ('difficulty',       'מספר אפסים מובילים הנדרשים ב-PoW'),
        ('chain',            'רשימת הבלוקים (ledger) בזיכרון'),
        ('balances',         'dict: public_key → יתרת AUR'),
        ('pending_txs',      'רשימת עסקות ממתינות'),
        ('gateway_client',   'RSA_Client ל-Gateway'),
        ('gateway_comm',     'Communication פעיל ל-Gateway'),
        ('node_dir',         'תיקיית אחסון מקומי node_<ip>_<port>/'),
        ('ledger_path',      'נתיב ledger.json'),
        ('balances_path',    'נתיב balances.json'),
    ],
    [
        ('start',                   '',                        'מתחבר ל-Gateway ומפעיל שרת עמיתים ולולאה ראשית'),
        ('connect_to_gateway',      '',                        'גילוי Gateway ב-UDP broadcast; מתחבר ושולח REGISTER_BLOCKCHAIN_NODE'),
        ('mine',                    'transaction: dict',       'PoW: מוצא nonce שמקיים hash(block).startswith("0"×difficulty); שומר לשרשרת; מחזיר block'),
        ('validate_tx',             'tx: dict',                'מאמת שיתרת השולח ≥ amount; מחזיר bool'),
        ('add_block',               'block: dict',             'מוסיף בלוק לשרשרת; מעדכן יתרות רק אם amount>0; שומר לדיסק'),
        ('get_balance',             'userpk',                  'מחזיר יתרת AUR לפי מפתח ציבורי (0.0 אם לא קיים)'),
        ('register_blockchain_node','',                        'שולח REGISTER_BLOCKCHAIN_NODE ל-Gateway עם chain_length'),
        ('notify_gateway',          'block',                   'שולח BROADCAST_TX_TO_VERIFY ל-Gateway'),
        ('notify_buy_success',      'tx_data',                 'שולח BUY_SUCCESS ל-Gateway'),
        ('notify_sell_success',     'tx_data',                 'שולח SELL_SUCCESS ל-Gateway'),
        ('send_balance',            'userpk',                  'שולח SEND_BALANCE ל-Gateway עם יתרת המפתח'),
        ('handle_tx_request_buy',   'msg',                     'כורה עסקת BUY; שולח BUY_SUCCESS ו-BROADCAST_TX_TO_VERIFY'),
        ('handle_tx_request_sell',  'msg',                     'כורה עסקת SELL; שולח SELL_SUCCESS ו-BROADCAST_TX_TO_VERIFY'),
        ('handle_create_balance',   'msg',                     'יוצר יתרת INITIAL_BALANCE למפתח חדש (רק אם לא קיים)'),
        ('_handle_mint_request',    'msg',                     'כורה ASSET_MINT (amount=0); שולח ASSET_SIGNED_IN_BLOCKCHAIN ל-Gateway'),
        ('_handle_unlist_request',  'msg',                     'כורה UNLIST_ASSET_FROM_BLOCKCHAIN (amount=0); שולח ASSET_UNLIST_SIGNED_IN_BLOCKCHAIN'),
        ('handle_broadcast_tx_to_verify','msg',                'מקבל בלוק מעמית; מוסיף לשרשרת אם תקין ורציף'),
        ('request_ledger_from_peer','peer_ip, peer_port',      'מתחבר לעמית, שולח GET_LEDGER, מקבל snapshot ומסנכרן מקומי'),
        ('request_balance_from_peer','peer_ip, peer_port, userpk','מבקש יתרה ספציפית מעמית ומעדכן balances מקומי'),
        ('handle_peer_connection',  'comm',                    'מטפל בחיבורי עמיתים: GET_LEDGER → snapshot; GET_BALANCE → יתרה'),
    ]
)

# ── Save ───────────────────────────────────────────────────────────────────────
OUT = r'C:\dev\aurex\aurex_modules_documentation.docx'
doc.save(OUT)
print(f'Saved: {OUT}')
